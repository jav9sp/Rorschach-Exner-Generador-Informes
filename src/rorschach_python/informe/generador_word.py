import pandas as pd
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def agregar_titulo(doc, texto, nivel=0):
    if nivel == 0:
        doc.add_heading(texto, level=1)
    else:
        doc.add_heading(texto, level=2)


def agregar_tabla_diccionario(doc, diccionario):
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = 'Light List'
    hdr_cells = tabla.rows[0].cells
    hdr_cells[0].text = 'Categoría'
    hdr_cells[1].text = 'Cantidad'
    for key, value in diccionario.items():
        row_cells = tabla.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)


def agregar_tabla_dataframe(doc, df):
    tabla = doc.add_table(rows=1, cols=len(df.columns))
    tabla.style = 'Light List'

    # Encabezados
    hdr_cells = tabla.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    # Filas de datos
    for _, row in df.iterrows():
        row_cells = tabla.add_row().cells
        es_fuera_de_rango = row['COMPARACIÓN'] != 'Dentro del rango'

        for i, cell_value in enumerate(row):
            cell = row_cells[i]
            cell.text = str(cell_value)

            # Si no está dentro del rango, se pinta
            if es_fuera_de_rango:
                shading_elm_1 = OxmlElement('w:shd')
                shading_elm_1.set(qn('w:fill'), 'ebba34')
                cell._tc.get_or_add_tcPr().append(shading_elm_1)


def agregar_interpretaciones(doc, titulo, interpretaciones):
    doc.add_heading(titulo, level=2)

    for texto in interpretaciones:
        if texto:
            parrafo = doc.add_paragraph(texto, style='Normal')


def generar_informe(nombre_archivo, secciones):
    doc = Document()
    doc.add_heading('Informe de Resultados Test de Rorschach', 0)

    for titulo, contenido in secciones.items():

        if isinstance(contenido, list):
            agregar_interpretaciones(doc, titulo, contenido)

        elif isinstance(contenido, dict):
            agregar_titulo(doc, titulo, nivel=1)
            agregar_tabla_diccionario(doc, contenido)

        elif isinstance(contenido, pd.DataFrame):
            agregar_titulo(doc, titulo, nivel=1)
            agregar_tabla_dataframe(doc, contenido)

    doc.save(nombre_archivo)
    print("\n¡Informe preliminar generado correctamente!\n")
